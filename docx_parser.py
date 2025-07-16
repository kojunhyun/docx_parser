# requirements: python-docx, pillow, pdfplumber, langchain-openai
import os
import re
import json
from PIL import Image
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
import subprocess
import pdfplumber

import base64
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import PydanticOutputParser
from pydantic import BaseModel, Field


from typing import Union, Any

import argparse

parser = argparse.ArgumentParser(description="DOCX 문서에서 텍스트, 표, 이미지 추출 및 JSON/Markdown 저장")

parser.add_argument("--data_path", type=str, required=True, help="입력 .docx 파일 경로")
parser.add_argument("--output_path", type=str, default="output", help="이미지 및 출력 마크다운 등을 저장할 디렉토리 (기본값: output)")

# 인자 파싱
args = parser.parse_args()

LocalLLM = ChatOpenAI(
    model="Qwen/Qwen2.5-VL-7B-Instruct",
    temperature=0, 
    api_key="EMPTY", 
    base_url="http://172.16.10.174:9901/v1",
    max_tokens=2048,
)


class ExplainImageExtractOCR(BaseModel):
    """Extract the following information from this image and return it exactly in the JSON"""
    explanation: str = Field(
        description="""summary image"""
    )
    ocr_text: str = Field(
        description="""Extract text"""
    )


class ExplainTable(BaseModel):
    """Extract the following information from this image and return it exactly in the JSON"""
    explanation: str = Field(
        description="""summary Table"""
    )


def analyze_table_with_llm(table_data, cell_images_info):
    prompt = """The following is a Markdown table. Based on the structure and content of the table, please provide a concise summary of what the table is about. 
    If there are any common patterns or characteristics among the entries, include them in the summary.
    If the table contains images, also describe the images and include them in the explanation.    
    Write the explanation in natural English, within 2–3 sentences. The response must be written in Korean.
    
    Table:
    {{markdown_table}}
    """

    new_markdown_table = table_to_markdown(table_data, cell_images_info)

    prompt_template = ChatPromptTemplate.from_messages([
        ("system", prompt),
        ("user", [            
            {"type": "text", "text": new_markdown_table}
        ])
    ])

    # Attach structured output
    EXPLAIN_LLM = LocalLLM.with_structured_output(ExplainTable)
    
    # Run the chain
    chain = prompt_template | EXPLAIN_LLM
    
    result = chain.invoke({"markdown_table": new_markdown_table})
    
    return result


def analyze_image_with_llm(image_path):
    prompt = """Please analyze the following image and return a response in this JSON format:    
    [
      "explanation": "<Brief description of the image. E.g., if it’s a document, mention the type; if it’s a photo, describe the scene.>",
      "ocr_text": "<Extract all readable text in the image as accurately as possible, including line breaks.>"
    ]
    
    - Follow this exact JSON format in your response.
    - Use 1–2 sentences for the explanation.
    - Please write the explanation in Korean only.
    - Ensure the OCR text is clean, readable, and complete.
    """
    
    with open(image_path, "rb") as f:
        encoded_image = base64.b64encode(f.read()).decode("utf-8")
    base64_image_url = f"data:image/jpeg;base64,{encoded_image}"

    prompt_template = ChatPromptTemplate.from_messages([
        ("system", prompt),
        ("user", [
            {"type": "image_url", "image_url": {"url": base64_image_url}},
            {"type": "text", "text": "Please extract the ExplainImageExtractOCR info and answer in JSON format only."}
        ])
    ])

    # Attach structured output
    OCR_LLM = LocalLLM.with_structured_output(ExplainImageExtractOCR)
    
    # Run the chain
    chain = prompt_template | OCR_LLM
    
    result = chain.invoke({"image": base64_image_url})
        
    return result


def convert_docx_to_pdf(docx_path, output_pdf_path):
    subprocess.run([
        "libreoffice",
        "--headless",
        "--convert-to", "pdf",
        docx_path,
        "--outdir", os.path.dirname(output_pdf_path)
    ], check=True)


def extract_page_text_map(pdf_path):
    page_text_map = []
    with pdfplumber.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            page_text_map.append({"page_number": i+1, "text": text})
    return page_text_map


def find_paragraph_page(paragraph_text, page_text_map, rendered_breaks=None, paragraph_index=None):
    # 1. Word 렌더링 정보가 있다면 우선 사용
    if rendered_breaks is not None and paragraph_index is not None:
        for (idx, page_num) in rendered_breaks:
            if paragraph_index <= idx:
                return page_num
        return rendered_breaks[-1][1] if rendered_breaks else 1

    # 2. fallback: PDF 기반
    for page in page_text_map:
        if paragraph_text[:30] in page["text"]:
            return page["page_number"]
    return None


def extract_rendered_breaks(doc):
    from lxml import etree
    body = doc.element.body
    rendered_breaks = []
    page_number = 1
    for idx, el in enumerate(body.iterchildren()):
        el_xml = etree.fromstring(el.xml.encode("utf-8"))
        if el_xml.xpath(".//w:lastRenderedPageBreak", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}):
            page_number += 1
        if el.tag.endswith("p"):
            rendered_breaks.append((idx, page_number))
    return rendered_breaks


def make_images_map(images_info):
    
    image_desc_map = {}

    for image_info in images_info:
        image_desc_map[image_info["image_path"]] = image_info["markdown"]

    return image_desc_map


def replace_images_in_markdown_table(markdown: str, image_desc_map: dict) -> str:
    """
    마크다운 테이블 내 ![](...) 형식의 이미지를 {설명} 으로 교체합니다.
    image_desc_map은 이미지 경로를 키로 하고 대체 텍스트를 값으로 가지는 dict입니다.
    """
    def replace_img(match):
        
        img_path = match.group(1).strip()
        alt_text = image_desc_map.get(img_path, "이미지")
        return f"{alt_text}"

    # 정규식으로 ![](path) 를 찾아 대체
    new_markdown = re.sub(r'!\[[^\]]*\]\(([^)]+)\)', replace_img, markdown)
        
    return new_markdown


def table_to_markdown(table_data, images_info=None):
    
    if not table_data:
        return ""
    header = table_data[0]
    rows = table_data[1:]
    
    md = "| " + " | ".join(header) + " |\n"
    md += "| " + " | " .join(["---"] * len(header)) + " |\n"

    for row in rows:
        md += "| " + " | ".join(row) + " |\n"
    
    if images_info is not None:
        if len(images_info) > 0:
            image_desc_map = make_images_map(images_info)
            cleaned_table = replace_images_in_markdown_table(md.strip(), image_desc_map)
            return cleaned_table
    
    return md.strip()


def extract_all_elements_in_order(docx_path, image_dir="images"):

    def para_processing(image_counter, paragraph_counter, table_counter):
    
        para = Paragraph(block, doc)
        text = para.text.strip()
        style = para.style.name.lower()
        page_num = find_paragraph_page(text, page_map, rendered_breaks, paragraph_counter)

        image_embeds = []
        for run in para.runs:
            for drawing in run._element.xpath(".//w:drawing"):
                for blip in drawing.xpath(".//a:blip"):
                    embed = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    if embed:
                        image_part = doc.part.related_parts[embed]
                        ext = image_part.content_type.split("/")[-1]
                        image_path = os.path.join(image_dir, f"image_{image_counter}.{ext}")
                        with open(image_path, "wb") as f:
                            f.write(image_part.blob)

                        image_result = analyze_image_with_llm(image_path)
                        
                        image_embeds.append({
                            "image_index": image_counter,
                            "image_path": image_path,
                            "location": f"paragraph_{paragraph_counter}",
                            "ocr_text": image_result.ocr_text,
                            "markdown": image_result.explanation,
                            "page_number": page_num
                        })
                        image_counter += 1

        if text or image_embeds:
            if "heading" in style or "제목" in style:
                try:
                    level = int(style.split()[-1])
                    md_text = "#" * level + " " + text
                except:
                    md_text = "# " + text
            else:
                md_text = text

            elements.append({
                "type": "paragraph",
                "paragraph_index": paragraph_counter,
                "page_number": page_num,
                "markdown": md_text,
                "images": image_embeds
            })
            paragraph_counter += 1

        return image_counter, paragraph_counter, table_counter

    
    def table_processing(image_counter, paragraph_counter, table_counter):

        table = Table(block, doc)
        table_data = []
        cell_images_info = []
        for r_idx, row in enumerate(table.rows):
            row_data = []
            for c_idx, cell in enumerate(row.cells):
                cell_text = " ".join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                cell_images = []
                for para in cell.paragraphs:
                    for run in para.runs:
                        for drawing in run._element.xpath(".//w:drawing"):
                            for blip in drawing.xpath(".//a:blip"):
                                embed = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                                if embed:
                                    image_part = doc.part.related_parts[embed]
                                    ext = image_part.content_type.split("/")[-1]
                                    image_path = os.path.join(image_dir, f"table_{table_counter}_img_{image_counter}.{ext}")
                                    with open(image_path, "wb") as f:
                                        f.write(image_part.blob)

                                    image_result = analyze_image_with_llm(image_path)
                                    
                                    cell_images.append({
                                        "image_index": image_counter,
                                        "image_path": image_path,
                                        "cell": [r_idx, c_idx],
                                        "ocr_text": image_result.ocr_text,
                                        "markdown": image_result.explanation,
                                    })
                                    image_counter += 1
                #row_data.append(cell_text + ("\n" + "\n".join([f"![img]({img['image_path']})" for img in cell_images]) if cell_images else ""))
                row_data.append(cell_text + ("\n".join([f"![img]({img['image_path']})" for img in cell_images]) if cell_images else ""))
                if cell_images:
                    cell_images_info.extend(cell_images)
            table_data.append(row_data)

        first_cell_text = table_data[0][0] if table_data and table_data[0] else ""
        page_num = find_paragraph_page(first_cell_text, page_map, None, None)

        explanation_table = analyze_table_with_llm(table_data, cell_images_info)
        #explanation_table = table_to_markdown(table_data, cell_images_info)

        elements.append({
            "type": "table",
            "table_index": table_counter,
            "page_number": page_num,
            "markdown": table_to_markdown(table_data),
            "explanation": explanation_table.explanation,
            "images": cell_images_info
        })
        table_counter += 1
        
        return image_counter, paragraph_counter, table_counter


    os.makedirs(image_dir, exist_ok=True)
    doc = Document(docx_path)

    pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
    convert_docx_to_pdf(docx_path, pdf_path)
    page_map = extract_page_text_map(pdf_path)
    rendered_breaks = extract_rendered_breaks(doc)

    elements = []
    image_counter = 0
    paragraph_counter = 0
    table_counter = 0

    for block in doc.element.body.iterchildren():
        tag = block.tag.split('}')[-1]

        if tag == 'p':
            
            image_counter, paragraph_counter, table_counter = para_processing(image_counter, paragraph_counter, table_counter)

        elif tag == 'tbl':
            
            image_counter, paragraph_counter, table_counter = table_processing(image_counter, paragraph_counter, table_counter)

    os.remove(pdf_path)

    return elements


if __name__ == '__main__':

    docx_path = args.data_path
    #docx_path = "/workspace/home/nuriadmin/pdf_loader_test/docs_parser/MIND AI2.docx"
    output = extract_all_elements_in_order(docx_path, image_dir=args.output_path)
    
    with open(os.path.join(args.output_path, "output.json"), "w", encoding="utf-8") as f:
        json.dump(output, f, indent=2, ensure_ascii=False)
